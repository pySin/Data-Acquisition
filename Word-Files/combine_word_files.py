## Find all Word(.docx) files in a folder and add all their texts together
## in 1 Word file. Tests are in 'test_w_combine.py'.

import docx
import os
import re

class integrate_f:

    def __init__(self):
        pass

    # Make a list of file names ending in '.docx'
    def filter_files(self, path, re_pattern):
        files_f = os.listdir(path)
        files_f = [x for x in files_f if re.match(re_pattern, x)]
        return files_f

    # Read and return the text from a '.docx' file.
    def read_docx(self, path, file_name):
        word_text = ''

        o_file = docx.Document(path+'\\'+file_name)
        for item in o_file.paragraphs:
            word_text = word_text+item.text

        return word_text

    # Combine the texts for all '.docx' files by calling the 'read_docx'
    # function for each file in list.
    def combine_texts(self, path, files_list):
        add_texts = ''

        for item in files_list:
            add_texts = add_texts+self.read_docx(path, item)+'\n'
        return add_texts

    # Create a new '.docx' document and fill it with the combined text.
    def make_docx(self, path, result_file, all_text):

        mydoc = docx.Document()

        mydoc.add_paragraph(all_text)
        mydoc.save(path+'\\'+result_file)

# Call all the functions in their proper order.
def call_functions():
    ob = integrate_f()

    path = 'C:\\Users\Sinan\Desktop\python\TaskFunctions\word_files'
    extension = 'docx'
    
    # Set the final, combined word file name.
    result_file = 'comb_word_f8.docx'

    re_pattern = '.*\.'+extension # Create a regular expression to match
                                  # anything except the chosen file extension.
    files_to_check = ob.filter_files(path, re_pattern)
    all_text = ob.combine_texts(path, files_to_check)
    ob.make_docx(path, result_file, all_text)

if __name__ == '__main__':
    call_functions()
